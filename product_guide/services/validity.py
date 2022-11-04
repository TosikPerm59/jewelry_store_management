import docx
import os


def check_outgoing_invoice(invoice_path):
    if not invoice_path.endswith('.docx'):
        return
    document = docx.Document(invoice_path)
    sender = 'ИП Александрова Елена Петровна "Golden_Sun", ИНН: 591792868890, улица Уральская дом 1'

    if 'Унифицированная форма № ТОРГ-12' in document.paragraphs[0].text:
        if len(document.paragraphs) == 3:
            return True


def check_path(path):
    if os.path.exists(path):
        return True


def check_file_path(file_path):
    if check_path(file_path):
        if os.path.isfile(file_path):
            return True


def check_id(_id):
    digits_count_list = [8, 10, 13]
    if not _id:
        return
    if len(str(_id)) in digits_count_list and str(_id).isdigit():
        return True
    else:
        return


def check_uin(uin):
    if not uin:
        return
    if len(str(uin)) == 16 and str(uin).isdigit():
        return True
    else:
        return


def isinteger(value):

    """ Функция проверяет элемент на причастность к целому числу.
    Возвращает boolean """

    try:
        int(value)
        return True
    except:
        return False


def isfloat(value):

    """ Функция проверяет элемент на причастность к вещественному числу.
        Возвращает boolean """

    try:
        float(value)
        return True
    except ValueError:
        return False
    except TypeError:
        return False


def check_word_exceptions(_string):

    """ Функция проверяет наличие слов исключений в передаваемой строке.
    Возвращает False в случае наличия слов исключений в строке или True при отсутствии """

    word_exceptions = ['585-й', '925-й', '0070', '585', '925', 'р-р']

    for exception in word_exceptions:
        if exception in _string:
            return False
    else:
        return True


def check_weight(weight):

    try:
        if isfloat(weight) and len(weight.split('.')[1]) <= 3 and float(weight) < 200:
            return True
    except IndexError:
        return False


