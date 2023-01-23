import os.path
import xlrd
import openpyxl
import warnings
import docx
from product_guide.services.request_classes import Request

warnings.simplefilter("ignore")


class ReadExcelFile:
    def __init__(self, file_handler_obj):
        Request.printCreateObject(self)
        self.file_handler_obj = file_handler_obj
        if file_handler_obj.file_extension == 'xls':
            self.rows_list, self.sheet = self.read_old_excel_file()
        else:
            self.rows_list, self.sheet = self.read_excel_file()

    def read_old_excel_file(self):
        print('Выполняется функция read_old_excel_file из класса ReadExcelFile')
        excel_file = xlrd.open_workbook(self.file_handler_obj.file_path)
        sheet = excel_file.sheet_by_index(0)
        rows_list = [sheet.row_values(rownum) for rownum in range(sheet.nrows)]
        for row in rows_list:
            for elem in row:
                if isinstance(elem, str):
                    elem_low = elem.replace('\n', '').lower() if '\n' in elem else elem.lower()
                    elem_ind = row.index(elem)
                    row.remove(elem)
                    row.insert(elem_ind, elem_low)
        return rows_list, sheet

    def read_excel_file(self):
        print('Выполняется функция read_excel_file из класса ReadExcelFile')
        excel_file = openpyxl.load_workbook(self.file_handler_obj.file_path)
        sheet = excel_file.active
        rows_list = [row for row in range(1, sheet.max_row + 1)]
        str_rows_list = self.create_string_rows_list(rows_list, sheet)
        return str_rows_list, sheet

    @staticmethod
    def create_string_rows_list(rows_list, sheet):
        cols = sheet.max_column
        str_rows_lst = []
        for row in rows_list:
            row_lst = []
            for col in range(cols):
                cell_text = sheet[row][col].value
                if isinstance(cell_text, str):
                    cell_text = cell_text.replace('\n', '').lower() if '\n' in cell_text else cell_text.lower()
                elif cell_text is None:
                    cell_text = ''
                row_lst.append(cell_text)
            str_rows_lst.append(row_lst)
        return str_rows_lst


class ReadWordFile:

    def __init__(self, file_handler_obj):
        Request.printCreateObject(self)
        self.file_handler_obj = file_handler_obj
        self.header_table, self.product_table = self.read_msword_file()

    def read_msword_file(self):
        document = docx.Document(self.file_handler_obj.file_path)
        header_table = document.tables[0]
        product_table = document.tables[1]
        return header_table, product_table



