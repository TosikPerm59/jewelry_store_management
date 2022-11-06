import os.path
import xlrd
import openpyxl
import warnings
import docx
warnings.simplefilter("ignore")


def read_excel_file(path_to_excel_file):
    path_to_excel_file = path_to_excel_file.replace('"', '') if '"' in path_to_excel_file else path_to_excel_file
    file_name = os.path.split(path_to_excel_file)[1]
    if file_name.endswith('.xls'):
        file_type = '.xls'
        excel_file = xlrd.open_workbook(path_to_excel_file)
        sheet = excel_file.sheet_by_index(0)
        rows_list = [sheet.row_values(rownum) for rownum in range(sheet.nrows)]
        for row in rows_list:
            for elem in row:
                if isinstance(elem, str):
                    elem_low = elem.lower()
                    elem_ind = row.index(elem)
                    row.remove(elem)
                    row.insert(elem_ind, elem_low)

        return rows_list, sheet, file_type

    elif file_name.endswith('.xlsx'):
        file_type = '.xlsx'
        excel_file = openpyxl.load_workbook(path_to_excel_file)
        sheet = excel_file.active
        rows_list = [row for row in range(1, sheet.max_row + 1)]
        return rows_list, sheet, file_type


def read_msword_file(path_to_msword_file):
    document = docx.Document(path_to_msword_file)
    header_table = document.tables[0]
    product_table = document.tables[1]
    return header_table, product_table
