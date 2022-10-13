from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import docx
import os
from product_guide.services.finders import find_weight


def change_outgoing_invoice(path):

    def set_cell_settings(cell_):
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].style.font.size = Pt(8)

    document = docx.Document(path)
    table = document.tables[1]
    max_row = len(table.rows)
    total_weight = 0

    for row in range(3, max_row - 21):
        string = table.rows[row].cells[4].text
        print(f'Поиск веса в строке ({string})')
        weight = find_weight(string.lower().split(' '))
        total_weight += float(weight)
        amount = (table.rows[row].cells[33].text.split(','))[0]

        cell = table.cell(row, 30)
        cell.text = amount
        set_cell_settings(cell)

        cell = table.cell(row, 33)
        cell.text = f'{weight} гр.'
        set_cell_settings(cell)

        cell = table.cell(row, 36)
        price = ''.join(((table.rows[row].cells[40].text.split(','))[0]).split('\xa0'))
        if price == '':
            price = ''.join(((table.rows[row].cells[35].text.split(','))[0]).split('\xa0'))
        cell.text = str(round(float(price) / int(amount) / float(weight), 2)).replace('.', ',')
        while len(cell.text.split(',')[1]) < 2:
            cell.text += '0'
        set_cell_settings(cell)

    cell = table.cell(max_row - 21, 33)
    cell.text = f'{str(round(total_weight, 2))} гр.'
    set_cell_settings(cell)

    document.save(path)
